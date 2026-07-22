from __future__ import annotations

import hashlib
import io
from collections.abc import Callable
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

CORPUS_ROOT = Path(__file__).parent / "corpus"


def parse_text(content: bytes) -> str:
    return _normalize(content.decode("utf-8", errors="strict"))


def parse_markdown(content: bytes) -> str:
    return parse_text(content)


def parse_html(content: bytes) -> str:
    markup = content.decode("utf-8", errors="strict")
    soup = BeautifulSoup(markup, "html.parser")
    for element in soup(["script", "style", "template", "noscript"]):
        element.decompose()
    return _normalize(soup.get_text("\n"))


def parse_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content), strict=True)
    if reader.is_encrypted:
        raise ValueError("encrypted-pdf")
    return _normalize("\n".join(page.extract_text() or "" for page in reader.pages))


def parse_docx(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return _normalize("\n".join(paragraphs))


def candidate_cases() -> dict[str, tuple[Callable[[bytes], str], bytes]]:
    return {
        "text-unicode": (parse_text, (CORPUS_ROOT / "text" / "unicode.txt").read_bytes()),
        "markdown-structure": (
            parse_markdown,
            (CORPUS_ROOT / "markdown" / "structure.md").read_bytes(),
        ),
        "html-active-content": (
            parse_html,
            (CORPUS_ROOT / "html" / "active-content.html").read_bytes(),
        ),
        "pdf-one-page": (parse_pdf, minimal_pdf()),
        "docx-headings": (parse_docx, minimal_docx()),
    }


def output_digest(value: str) -> str:
    return "sha256:" + hashlib.sha256(value.encode("utf-8")).hexdigest()


def minimal_docx() -> bytes:
    document = Document()
    document.add_heading("Apvero parser benchmark", level=1)
    document.add_paragraph("Deterministic DOCX paragraph.")
    target = io.BytesIO()
    document.save(target)
    return target.getvalue()


def minimal_pdf() -> bytes:
    text = "Apvero parser benchmark page one."
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{index} 0 obj\n".encode("ascii"))
        result.extend(body)
        result.extend(b"\nendobj\n")
    xref = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    result.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return bytes(result)


def _normalize(value: str) -> str:
    lines = [line.strip() for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()
