from __future__ import annotations

import hashlib
import io
import zipfile
from uuid import UUID, uuid4

import pytest
from docx import Document
from pypdf import PdfWriter

from apvero_worker.document_processing import (
    PROCESSING_PROFILE,
    WorkerProcessingError,
    process_document_snapshot,
)
from benchmarks.parser_candidates import minimal_docx, minimal_pdf


def test_all_five_media_types_are_deterministic_and_traceable() -> None:
    cases = {
        "text/plain": "First line\nSecond 😀 line".encode(),
        "text/markdown": b"# Policy\n\nApproved amount: 100.",
        "text/html": b"<h1>Policy</h1><p>Retained</p><script>discarded()</script>",
        "application/pdf": minimal_pdf(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": (
            minimal_docx()
        ),
    }

    for media_type, content in cases.items():
        request_id = uuid4()
        revision_id = uuid4()
        first = _process(request_id, revision_id, media_type, content)
        second = _process(request_id, revision_id, media_type, content)

        assert first == second
        assert first.processing_profile == PROCESSING_PROFILE
        assert first.content_digest == _digest(content)
        assert first.parser_version.startswith("apvero-")
        assert first.chunker_version == "apvero-boundary@1.0.0"
        assert first.documents
        assert all(document.chunks for document in first.documents)
        assert [document.ordinal for document in first.documents] == list(
            range(len(first.documents))
        )
        for document in first.documents:
            assert [chunk.ordinal for chunk in document.chunks] == list(
                range(len(document.chunks))
            )
            for chunk in document.chunks:
                assert chunk.content_digest == _digest(chunk.text.encode())
                assert chunk.end_offset > chunk.start_offset


def test_unicode_offsets_and_markdown_anchors_use_code_points() -> None:
    content = "# 标题\nA😀B\n下一行".encode()
    result = _process(uuid4(), uuid4(), "text/markdown", content)

    document = result.documents[0]
    chunk = document.chunks[0]
    assert chunk.start_offset == 0
    assert chunk.end_offset == len("# 标题\nA😀B\n下一行")
    assert chunk.anchors.heading == "标题"
    assert chunk.anchors.paragraph == 1
    assert chunk.anchors.line_start == 1
    assert chunk.anchors.line_end == 3


def test_html_removes_executable_content_and_keeps_heading_lineage() -> None:
    result = _process(
        uuid4(),
        uuid4(),
        "text/html",
        b"<h2>Security</h2><p>Keep this.</p><script>remove-secret()</script>",
    )

    chunk = result.documents[0].chunks[0]
    assert "Keep this." in chunk.text
    assert "remove-secret" not in chunk.text
    assert chunk.anchors.heading == "Security"


def test_pdf_and_docx_keep_page_and_paragraph_anchors() -> None:
    pdf = _process(uuid4(), uuid4(), "application/pdf", minimal_pdf())
    assert pdf.documents[0].chunks[0].anchors.page == 1

    document = Document()
    document.add_heading("Operations", level=1)
    document.add_paragraph("Restart safely.")
    target = io.BytesIO()
    document.save(target)
    docx = _process(
        uuid4(),
        uuid4(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        target.getvalue(),
    )
    assert docx.documents[0].title == "Operations"
    assert docx.documents[0].chunks[0].anchors.paragraph == 1


def test_digest_mismatch_malformed_encrypted_and_archive_bomb_fail_closed() -> None:
    request_id = uuid4()
    with pytest.raises(WorkerProcessingError, match="WORKER_CONTENT_DIGEST_MISMATCH"):
        process_document_snapshot(
            request_id=request_id,
            source_revision_id=uuid4(),
            content_digest=_digest(b"different"),
            media_type="text/plain",
            processing_profile=PROCESSING_PROFILE,
            content=b"actual",
        )

    with pytest.raises(WorkerProcessingError, match="WORKER_MALFORMED_DOCUMENT"):
        _process(request_id, uuid4(), "application/pdf", b"%PDF-malformed")

    encrypted = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    writer.write(encrypted)
    with pytest.raises(WorkerProcessingError, match="WORKER_ENCRYPTED_DOCUMENT"):
        _process(request_id, uuid4(), "application/pdf", encrypted.getvalue())

    bomb = io.BytesIO()
    with zipfile.ZipFile(bomb, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<document/>")
        archive.writestr("word/bomb.bin", b"0" * 1_000_000)
    with pytest.raises(WorkerProcessingError, match="WORKER_ARCHIVE_EXPANSION_LIMIT"):
        _process(
            request_id,
            uuid4(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            bomb.getvalue(),
        )


def test_processing_deadline_and_profile_are_enforced() -> None:
    content = b"bounded"
    with pytest.raises(WorkerProcessingError, match="WORKER_PROCESSING_TIMEOUT"):
        process_document_snapshot(
            request_id=uuid4(),
            source_revision_id=uuid4(),
            content_digest=_digest(content),
            media_type="text/plain",
            processing_profile=PROCESSING_PROFILE,
            content=content,
            timeout_seconds=0,
        )
    with pytest.raises(WorkerProcessingError, match="WORKER_INVALID_REQUEST"):
        process_document_snapshot(
            request_id=uuid4(),
            source_revision_id=uuid4(),
            content_digest=_digest(content),
            media_type="text/plain",
            processing_profile="mutable-latest",
            content=content,
        )


def test_overlapping_chunks_reconstruct_the_normalized_document() -> None:
    content = (("segment 😀 " * 900) + "finished").encode()
    result = _process(uuid4(), uuid4(), "text/plain", content)
    document = result.documents[0]

    assert len(document.chunks) > 1
    reconstructed = document.chunks[0].text
    for chunk in document.chunks[1:]:
        overlap = len(reconstructed) - chunk.start_offset
        assert reconstructed[-overlap:] == chunk.text[:overlap]
        reconstructed += chunk.text[overlap:]
    assert _digest(reconstructed.encode()) == document.content_digest


def _process(
    request_id: UUID, source_revision_id: UUID, media_type: str, content: bytes
):
    return process_document_snapshot(
        request_id=request_id,
        source_revision_id=source_revision_id,
        content_digest=_digest(content),
        media_type=media_type,
        processing_profile=PROCESSING_PROFILE,
        content=content,
    )


def _digest(value: bytes) -> str:
    return "sha256:" + hashlib.sha256(value).hexdigest()
