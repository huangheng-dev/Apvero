from __future__ import annotations

import hashlib
import io
import re
import time
import unicodedata
import zipfile
from dataclasses import dataclass
from typing import Final
from uuid import UUID

from bs4 import BeautifulSoup, Tag
from docx import Document
from pypdf import PdfReader

from apvero_worker.models import (
    ProcessedChunk,
    ProcessedDocument,
    ProcessedDocumentBatch,
    SourceAnchors,
)

PROCESSING_PROFILE: Final = "apvero-default@1.0.0"
CHUNKER_VERSION: Final = "apvero-boundary@1.0.0"
MAX_INPUT_BYTES: Final = 5_242_880
MAX_OUTPUT_CODE_POINTS: Final = 10_000_000
MAX_PAGES: Final = 500
MAX_DOCX_ENTRIES: Final = 2_048
MAX_DOCX_EXPANDED_BYTES: Final = 20_971_520
MAX_DOCX_ENTRY_RATIO: Final = 100
MAX_DOCUMENTS: Final = 10_000
MAX_CHUNKS: Final = 100_000
CHUNK_SIZE: Final = 4_000
CHUNK_OVERLAP: Final = 200
PROCESSING_TIMEOUT_SECONDS: Final = 10.0

MEDIA_TYPES: Final = {
    "text/plain",
    "text/markdown",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/html",
}

PARSER_VERSIONS: Final = {
    "text/plain": "apvero-text@1.0.0",
    "text/markdown": "apvero-markdown@1.0.0",
    "application/pdf": "apvero-pdf@1.0.0",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": (
        "apvero-docx@1.0.0"
    ),
    "text/html": "apvero-html@1.0.0",
}


class WorkerProcessingError(Exception):
    def __init__(
        self,
        code: str,
        *,
        status: int,
        retryable: bool = False,
        request_id: UUID | None = None,
    ) -> None:
        super().__init__(code)
        self.code = code
        self.status = status
        self.retryable = retryable
        self.request_id = request_id


@dataclass(frozen=True)
class _AnchorSpan:
    start: int
    end: int
    page: int | None = None
    heading: str | None = None
    paragraph: int | None = None


@dataclass(frozen=True)
class _ParsedDocument:
    title: str | None
    text: str
    spans: tuple[_AnchorSpan, ...]


def process_document_snapshot(
    *,
    request_id: UUID,
    source_revision_id: UUID,
    content_digest: str,
    media_type: str,
    processing_profile: str,
    content: bytes,
    timeout_seconds: float = PROCESSING_TIMEOUT_SECONDS,
) -> ProcessedDocumentBatch:
    started = time.monotonic()
    _validate_request(
        request_id=request_id,
        content_digest=content_digest,
        media_type=media_type,
        processing_profile=processing_profile,
        content=content,
    )
    _check_deadline(started, timeout_seconds, request_id)
    try:
        parsed = _parse(media_type, content, request_id, started, timeout_seconds)
    except WorkerProcessingError as exception:
        if exception.request_id is None:
            exception.request_id = request_id
        raise
    except (UnicodeDecodeError, ValueError, OSError, zipfile.BadZipFile) as exception:
        raise WorkerProcessingError(
            "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
        ) from exception
    _check_deadline(started, timeout_seconds, request_id)
    if not parsed or len(parsed) > MAX_DOCUMENTS:
        raise WorkerProcessingError(
            "WORKER_RESOURCE_EXHAUSTED", status=503, retryable=True, request_id=request_id
        )

    documents: list[ProcessedDocument] = []
    total_chunks = 0
    for ordinal, document in enumerate(parsed):
        _check_deadline(started, timeout_seconds, request_id)
        chunks = _chunk(document)
        total_chunks += len(chunks)
        if total_chunks > MAX_CHUNKS:
            raise WorkerProcessingError(
                "WORKER_RESOURCE_EXHAUSTED", status=503, retryable=True, request_id=request_id
            )
        documents.append(
            ProcessedDocument(
                ordinal=ordinal,
                title=document.title,
                content_digest=_digest(document.text.encode("utf-8")),
                chunks=chunks,
            )
        )

    return ProcessedDocumentBatch(
        request_id=str(request_id),
        source_revision_id=str(source_revision_id),
        content_digest=content_digest,
        processing_profile=processing_profile,
        parser_version=PARSER_VERSIONS[media_type],
        chunker_version=CHUNKER_VERSION,
        documents=documents,
        warnings=[],
    )


def _validate_request(
    *,
    request_id: UUID,
    content_digest: str,
    media_type: str,
    processing_profile: str,
    content: bytes,
) -> None:
    if processing_profile != PROCESSING_PROFILE:
        raise WorkerProcessingError("WORKER_INVALID_REQUEST", status=400, request_id=request_id)
    if media_type not in MEDIA_TYPES:
        raise WorkerProcessingError(
            "WORKER_UNSUPPORTED_MEDIA_TYPE", status=415, request_id=request_id
        )
    if not content or len(content) > MAX_INPUT_BYTES:
        raise WorkerProcessingError(
            "WORKER_CONTENT_TOO_LARGE", status=413, request_id=request_id
        )
    if not re.fullmatch(r"sha256:[a-f0-9]{64}", content_digest):
        raise WorkerProcessingError("WORKER_INVALID_REQUEST", status=400, request_id=request_id)
    if _digest(content) != content_digest:
        raise WorkerProcessingError(
            "WORKER_CONTENT_DIGEST_MISMATCH", status=400, request_id=request_id
        )


def _parse(
    media_type: str,
    content: bytes,
    request_id: UUID,
    started: float,
    timeout_seconds: float,
) -> tuple[_ParsedDocument, ...]:
    if media_type == "text/plain":
        return (_text_document(content),)
    if media_type == "text/markdown":
        return (_markdown_document(content),)
    if media_type == "text/html":
        return (_html_document(content),)
    if media_type == "application/pdf":
        return _pdf_documents(content, request_id, started, timeout_seconds)
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return (_docx_document(content, request_id),)
    raise WorkerProcessingError(
        "WORKER_UNSUPPORTED_MEDIA_TYPE", status=415, request_id=request_id
    )


def _text_document(content: bytes) -> _ParsedDocument:
    text = _normalize(content.decode("utf-8", errors="strict"))
    _require_text(text)
    return _ParsedDocument(title=None, text=text, spans=())


def _markdown_document(content: bytes) -> _ParsedDocument:
    text = _normalize(content.decode("utf-8", errors="strict"))
    _require_text(text)
    spans: list[_AnchorSpan] = []
    heading: str | None = None
    offset = 0
    paragraph = 0
    for line in text.splitlines(keepends=True):
        value = line.rstrip("\n")
        match = re.match(r"^#{1,6}\s+(.+?)\s*$", value)
        if match:
            heading = match.group(1)[:1_000]
        if value.strip():
            paragraph += 1
            spans.append(
                _AnchorSpan(
                    offset,
                    offset + len(value),
                    heading=heading,
                    paragraph=paragraph,
                )
            )
        offset += len(line)
    title = next((span.heading for span in spans if span.heading), None)
    return _ParsedDocument(title=title, text=text, spans=tuple(spans))


def _html_document(content: bytes) -> _ParsedDocument:
    markup = content.decode("utf-8", errors="strict")
    soup = BeautifulSoup(markup, "html.parser")
    for element in soup(["script", "style", "template", "noscript", "object", "embed", "iframe"]):
        element.decompose()
    root = soup.body or soup
    blocks: list[tuple[str, str | None, int]] = []
    heading: str | None = None
    paragraph = 0
    for element in root.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"], recursive=True):
        if not isinstance(element, Tag) or element.find_parent(["p", "li"]):
            continue
        value = _normalize(element.get_text(" ", strip=True))
        if not value:
            continue
        if element.name and element.name.startswith("h"):
            heading = value[:1_000]
        paragraph += 1
        blocks.append((value, heading, paragraph))
    if not blocks:
        fallback = _normalize(root.get_text("\n", strip=True))
        _require_text(fallback)
        blocks.append((fallback, None, 1))
    return _document_from_blocks(blocks, title=next((item[1] for item in blocks if item[1]), None))


def _pdf_documents(
    content: bytes,
    request_id: UUID,
    started: float,
    timeout_seconds: float,
) -> tuple[_ParsedDocument, ...]:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True)
    except Exception as exception:
        raise WorkerProcessingError(
            "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
        ) from exception
    if reader.is_encrypted:
        raise WorkerProcessingError(
            "WORKER_ENCRYPTED_DOCUMENT", status=422, request_id=request_id
        )
    if len(reader.pages) > MAX_PAGES:
        raise WorkerProcessingError("WORKER_PAGE_LIMIT", status=422, request_id=request_id)
    documents: list[_ParsedDocument] = []
    for page_number, page in enumerate(reader.pages, start=1):
        _check_deadline(started, timeout_seconds, request_id)
        try:
            text = _normalize(page.extract_text() or "")
        except Exception as exception:
            raise WorkerProcessingError(
                "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
            ) from exception
        if not text:
            continue
        documents.append(
            _ParsedDocument(
                title=f"Page {page_number}",
                text=text,
                spans=(_AnchorSpan(0, len(text), page=page_number),),
            )
        )
    if not documents:
        raise WorkerProcessingError(
            "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
        )
    return tuple(documents)


def _docx_document(content: bytes, request_id: UUID) -> _ParsedDocument:
    _inspect_docx_archive(content, request_id)
    try:
        document = Document(io.BytesIO(content))
    except Exception as exception:
        raise WorkerProcessingError(
            "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
        ) from exception
    blocks: list[tuple[str, str | None, int]] = []
    heading: str | None = None
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        value = _normalize(paragraph.text)
        if not value:
            continue
        style = paragraph.style.name if paragraph.style is not None else ""
        if style.lower().startswith("heading"):
            heading = value[:1_000]
        blocks.append((value, heading, paragraph_number))
    if not blocks:
        raise WorkerProcessingError(
            "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
        )
    return _document_from_blocks(blocks, title=next((item[1] for item in blocks if item[1]), None))


def _inspect_docx_archive(content: bytes, request_id: UUID) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise WorkerProcessingError(
                    "WORKER_ARCHIVE_EXPANSION_LIMIT", status=422, request_id=request_id
                )
            names = {entry.filename.lower() for entry in entries}
            if "[content_types].xml" not in names or "word/document.xml" not in names:
                raise WorkerProcessingError(
                    "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
                )
            if any(name.endswith("vbaproject.bin") for name in names):
                raise WorkerProcessingError(
                    "WORKER_UNSUPPORTED_MEDIA_TYPE", status=415, request_id=request_id
                )
            expanded = 0
            for entry in entries:
                if entry.flag_bits & 0x1:
                    raise WorkerProcessingError(
                        "WORKER_ENCRYPTED_DOCUMENT", status=422, request_id=request_id
                    )
                expanded += entry.file_size
                ratio = entry.file_size / max(entry.compress_size, 1)
                if expanded > MAX_DOCX_EXPANDED_BYTES or ratio > MAX_DOCX_ENTRY_RATIO:
                    raise WorkerProcessingError(
                        "WORKER_ARCHIVE_EXPANSION_LIMIT", status=422, request_id=request_id
                    )
    except WorkerProcessingError:
        raise
    except zipfile.BadZipFile as exception:
        raise WorkerProcessingError(
            "WORKER_MALFORMED_DOCUMENT", status=422, request_id=request_id
        ) from exception


def _document_from_blocks(
    blocks: list[tuple[str, str | None, int]], *, title: str | None
) -> _ParsedDocument:
    parts: list[str] = []
    spans: list[_AnchorSpan] = []
    offset = 0
    for value, heading, paragraph in blocks:
        if parts:
            parts.append("\n")
            offset += 1
        start = offset
        parts.append(value)
        offset += len(value)
        spans.append(_AnchorSpan(start, offset, heading=heading, paragraph=paragraph))
    text = "".join(parts)
    _require_text(text)
    return _ParsedDocument(title=title, text=text, spans=tuple(spans))


def _chunk(document: _ParsedDocument) -> list[ProcessedChunk]:
    chunks: list[ProcessedChunk] = []
    start = 0
    while start < len(document.text):
        hard_end = min(start + CHUNK_SIZE, len(document.text))
        end = _nearest_boundary(document.text, start, hard_end)
        if end <= start:
            end = hard_end
        value = document.text[start:end]
        span = _span_at(document.spans, start)
        line_start = document.text.count("\n", 0, start) + 1
        line_end = document.text.count("\n", 0, max(end - 1, start)) + 1
        chunks.append(
            ProcessedChunk(
                ordinal=len(chunks),
                text=value,
                content_digest=_digest(value.encode("utf-8")),
                start_offset=start,
                end_offset=end,
                anchors=SourceAnchors(
                    page=span.page if span else None,
                    heading=span.heading if span else None,
                    paragraph=span.paragraph if span else None,
                    line_start=line_start,
                    line_end=line_end,
                ),
            )
        )
        if end == len(document.text):
            break
        start = max(start + 1, end - CHUNK_OVERLAP)
    return chunks


def _span_at(spans: tuple[_AnchorSpan, ...], offset: int) -> _AnchorSpan | None:
    return next((span for span in spans if span.start <= offset < span.end), None)


def _nearest_boundary(text: str, start: int, hard_end: int) -> int:
    if hard_end == len(text):
        return hard_end
    minimum = start + ((hard_end - start) * 3 // 5)
    for delimiter in ("\n\n", "\n", ". ", "。", "! ", "！", "? ", "？"):
        boundary = text.rfind(delimiter, minimum, hard_end)
        if boundary >= minimum:
            return boundary + len(delimiter)
    return hard_end


def _normalize(value: str) -> str:
    normalized = unicodedata.normalize("NFC", value.replace("\r\n", "\n").replace("\r", "\n"))
    sanitized = "".join(
        character
        for character in normalized
        if character in {"\n", "\t"} or unicodedata.category(character) != "Cc"
    )
    if len(sanitized) > MAX_OUTPUT_CODE_POINTS:
        raise WorkerProcessingError("WORKER_RESOURCE_EXHAUSTED", status=503, retryable=True)
    return "\n".join(line.rstrip() for line in sanitized.split("\n")).strip()


def _require_text(value: str) -> None:
    if not value:
        raise WorkerProcessingError("WORKER_MALFORMED_DOCUMENT", status=422)


def _check_deadline(started: float, timeout_seconds: float, request_id: UUID) -> None:
    if timeout_seconds <= 0 or time.monotonic() - started > timeout_seconds:
        raise WorkerProcessingError(
            "WORKER_PROCESSING_TIMEOUT", status=503, retryable=True, request_id=request_id
        )


def _digest(value: bytes) -> str:
    return "sha256:" + hashlib.sha256(value).hexdigest()
